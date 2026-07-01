from rest_framework import generics, status, views
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework_simplejwt.views import TokenObtainPairView
from django.contrib.auth import get_user_model
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
import json
import re
import random
import os
import threading
from django.utils import timezone

from . import services

from .models import UserProfile, Assessment, Result, AssessmentSession, ProctoringViolation, CVAnalysis, CodingSubmission, AIInterviewSession, SkillGapAnalysis, TestAssignment, TestInvitation
from .serializers import (
    UserSerializer, 
    RegisterSerializer, 
    UserProfileUpdateSerializer,
    AssessmentSerializer,
    ResultSerializer,
    AssessmentSessionSerializer,
    ProctoringViolationSerializer,
    CVAnalysisSerializer,
    CodingSubmissionSerializer,
    AIInterviewSessionSerializer,
    SkillGapAnalysisSerializer,
    TestAssignmentSerializer,
    TestInvitationSerializer
)

User = get_user_model()

import requests as http_requests

TURNSTILE_SECRET_KEY = os.environ.get('TURNSTILE_SECRET_KEY', '0x4AAAAAACyCFsfphnqqQoBZoa9fgjtIqwA')
TURNSTILE_VERIFY_URL = 'https://challenges.cloudflare.com/turnstile/v0/siteverify'

class TurnstileLoginView(TokenObtainPairView):
    """Custom login view that verifies Cloudflare Turnstile before issuing JWT tokens.
    
    Turnstile verification is enforced when a turnstile_token is present in the 
    request data (i.e., from the login page). When no token is present (e.g., 
    auto-login after registration), the view falls through to normal JWT auth.
    """
    
    def post(self, request, *args, **kwargs):
        turnstile_token = request.data.get('turnstile_token')
        
        # Only verify if a token was submitted (login page always includes one)
        if turnstile_token:
            try:
                verify_response = http_requests.post(TURNSTILE_VERIFY_URL, data={
                    'secret': TURNSTILE_SECRET_KEY,
                    'response': turnstile_token,
                }, timeout=10)
                
                result = verify_response.json()
                
                if not result.get('success'):
                    return Response(
                        {'turnstile': 'Captcha verification failed. Please try again.'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except Exception:
                # If Cloudflare is unreachable, allow login anyway (graceful degradation)
                pass
        
        # Proceed with normal JWT token creation
        return super().post(request, *args, **kwargs)

class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    permission_classes = (AllowAny,)
    serializer_class = RegisterSerializer

class ForgotPasswordView(views.APIView):
    """Accept an email, validate the user exists, and return a password reset uid+token."""
    permission_classes = (AllowAny,)

    def post(self, request):
        email = request.data.get('email', '').strip()
        if not email:
            return Response({'error': 'Email is required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return Response({'error': 'No account found with this email address.'}, status=status.HTTP_404_NOT_FOUND)

        # Generate token using Django's built-in system (tied to user's password hash)
        uid = urlsafe_base64_encode(force_bytes(user.pk))
        token = default_token_generator.make_token(user)

        from django.core.mail import send_mail
        from django.conf import settings
        
        frontend_url = request.headers.get('origin', 'http://localhost:5173')
        reset_link = f"{frontend_url}/reset-password?uid={uid}&token={token}"
        
        subject = "Password Reset Request"
        message = f"Hello,\n\nWe received a request to reset the password for your account.\nPlease click the link below to set a new password:\n\n{reset_link}\n\nIf you did not request a password reset, please ignore this email."
        
        try:
            send_mail(
                subject,
                message,
                settings.EMAIL_HOST_USER,
                [email],
                fail_silently=False,
            )
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({'error': 'Failed to send reset email. Please try again later.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response({
            'message': 'Password reset link sent to your email.',
        })

class ResetPasswordView(views.APIView):
    """Accept uid, token, and new_password to reset the user's password."""
    permission_classes = (AllowAny,)

    def post(self, request):
        uid = request.data.get('uid', '')
        token = request.data.get('token', '')
        new_password = request.data.get('new_password', '')

        if not uid or not token or not new_password:
            return Response({'error': 'All fields are required.'}, status=status.HTTP_400_BAD_REQUEST)

        if len(new_password) < 8:
            return Response({'error': 'Password must be at least 8 characters.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            user_pk = force_str(urlsafe_base64_decode(uid))
            user = User.objects.get(pk=user_pk)
        except (ValueError, TypeError, User.DoesNotExist):
            return Response({'error': 'Invalid reset link.'}, status=status.HTTP_400_BAD_REQUEST)

        if not default_token_generator.check_token(user, token):
            return Response({'error': 'Reset token has expired or is invalid.'}, status=status.HTTP_400_BAD_REQUEST)

        user.set_password(new_password)
        user.save()

        return Response({'message': 'Password has been reset successfully.'})

class UserProfileView(generics.RetrieveUpdateAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = UserProfileUpdateSerializer
    parser_classes = (MultiPartParser, FormParser)

    def get_object(self):
        # We always return the logged-in user's profile
        profile, created = UserProfile.objects.get_or_create(user=self.request.user)
        return profile
    
    def retrieve(self, request, *args, **kwargs):
        # Also include user details when fetching profile
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        user_serializer = UserSerializer(request.user)
        
        return Response({
            'user': user_serializer.data,
            'profile': serializer.data
        })

class CVProfileExtractView(views.APIView):
    permission_classes = (IsAuthenticated,)
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('resume')
        if not file_obj:
            return Response({"error": "No resume file provided"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # We reuse the extraction logic from CVTestGenerateView
            import PyPDF2
            import docx

            file_ext = os.path.splitext(file_obj.name)[1].lower()
            text = ""

            if file_ext == '.pdf':
                pdf_reader = PyPDF2.PdfReader(file_obj)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            elif file_ext in ['.doc', '.docx']:
                doc = docx.Document(file_obj)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:
                return Response({"error": "Unsupported file format. Please upload PDF, DOC, or DOCX."}, status=status.HTTP_400_BAD_REQUEST)

            if len(text.strip()) < 50:
                return Response({"error": "Could not extract enough text from the file."}, status=status.HTTP_400_BAD_REQUEST)

            # Call AI Service
            profile_data = services.extract_profile_from_cv(text)

            if not profile_data:
                return Response({"error": "Failed to extract profile from CV."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response(profile_data)

        except Exception as e:
            return Response({"error": f"Error processing file: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ATSResumeCheckView(views.APIView):
    permission_classes = (IsAuthenticated,)
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        use_existing = request.data.get('use_existing_cv') == 'true' or request.data.get('use_existing_cv') is True
        file_obj = request.FILES.get('resume')
        
        text = ""
        
        if use_existing:
            profile = request.user.profile
            if not profile.resume:
                return Response({"error": "No existing resume found in profile."}, status=status.HTTP_400_BAD_REQUEST)
            
            # Open existing file
            file_path = profile.resume.path
            if not os.path.exists(file_path):
                return Response({
                    "error": "The resume file associated with your profile could not be found on our servers. Please re-upload your resume in the 'Upload New' tab to continue."
                }, status=status.HTTP_404_NOT_FOUND)
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            try:
                import PyPDF2
                import docx
                if file_ext == '.pdf':
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                elif file_ext in ['.doc', '.docx']:
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
            except Exception as e:
                return Response({"error": f"Error reading existing resume: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        elif file_obj:
            try:
                import PyPDF2
                import docx
                file_ext = os.path.splitext(file_obj.name)[1].lower()
                if file_ext == '.pdf':
                    pdf_reader = PyPDF2.PdfReader(file_obj)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                elif file_ext in ['.doc', '.docx']:
                    doc = docx.Document(file_obj)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                else:
                    return Response({"error": "Unsupported file format."}, status=status.HTTP_400_BAD_REQUEST)
            except Exception as e:
                return Response({"error": f"Error processing uploaded resume: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            return Response({"error": "No resume source provided (file or use_existing_cv)."}, status=status.HTTP_400_BAD_REQUEST)

        if len(text.strip()) < 50:
            return Response({"error": "Could not extract enough text from the resume."}, status=status.HTTP_400_BAD_REQUEST)

        # Call AI Service
        ats_report = services.analyze_resume_ats(text)

        if not ats_report:
            return Response({"error": "Failed to analyze resume with AI."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response(ats_report)

class ConfidenceCheckView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        transcript = request.data.get('transcript')
        duration = request.data.get('duration') # in seconds
        
        if not transcript:
            return Response({"error": "No transcript provided."}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            duration = float(duration)
        except (TypeError, ValueError):
            return Response({"error": "Invalid duration."}, status=status.HTTP_400_BAD_REQUEST)

        # Call AI Service
        report = services.analyze_voice_confidence(transcript, duration)

        if not report:
            return Response({"error": "Failed to analyze voice confidence."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response(report)

class AIMentorChatView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        message = request.data.get('message')
        if not message:
            return Response({"error": "No message provided."}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get personalized response
        ai_response = services.get_ai_mentor_response(request.user, message)
        
        return Response({"response": ai_response})

class RecentActivityView(generics.ListAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = ResultSerializer

    def get_queryset(self):
        return Result.objects.filter(user=self.request.user).order_by('-completed_at')[:5]

class DashboardStatsView(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        user = request.user
        
        # Use the new aggregated service
        stats = services.get_aggregated_user_stats(user)
        ai_insights = services.generate_dashboard_insights(user)
        
        # Additional legacy fields for UI compatibility if needed
        coding_solved = CodingSubmission.objects.filter(user=user, status='Accepted').count()
        sessions = AssessmentSession.objects.filter(user=user)
        secure_exams_count = sessions.filter(status='completed').count()
        
        total_violations = ProctoringViolation.objects.filter(session__user=user).count()
        integrity_score = max(0, 100 - (total_violations * 5))
        
        return Response({
            'total_tests': stats['total_tests'],
            'total_interviews': stats['total_interviews'],
            'total_assessments': stats['total_tests'] + stats['total_interviews'],
            'average_score': stats['avg_score'],
            'coding_solved': coding_solved,
            'secure_exams_count': secure_exams_count,
            'integrity_score': f"{integrity_score}%",
            'top_skill': stats['top_skill'],
            'weakest_skill': stats['weakest_skill'],
            'ai_insights': ai_insights
        })

class AssessmentListView(generics.ListAPIView):
    queryset = Assessment.objects.all()
    serializer_class = AssessmentSerializer
    permission_classes = (IsAuthenticated,)
    filterset_fields = ['category', 'difficulty']
    search_fields = ['title', 'description']
    ordering_fields = ['enrolled_count', 'estimated_time', 'id']
    ordering = ['-enrolled_count'] # Default to popular

class AssessmentMetadataView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        categories = Assessment.objects.values_list('category', flat=True).distinct()
        difficulties = Assessment.objects.values_list('difficulty', flat=True).distinct()
        
        return Response({
            'categories': sorted(list(set(categories))),
            'difficulties': sorted(list(set(difficulties)))
        })

class AssessmentSessionStartView(generics.CreateAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = AssessmentSessionSerializer

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class ProctoringViolationLogView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, session_id):
        # session_id can be for AssessmentSession or we check for AIInterviewSession
        is_interview = request.data.get('is_interview', False)
        try:
            violation_type = request.data.get('type')
            strike_number = request.data.get('strike_number')
            
            if is_interview:
                session = AIInterviewSession.objects.get(id=session_id, user=request.user)
                violation = ProctoringViolation.objects.create(
                    interview_session=session,
                    type=violation_type,
                    strike_number=strike_number,
                    comment=request.data.get('comment', '')
                )
            else:
                session = AssessmentSession.objects.get(id=session_id, user=request.user)
                violation = ProctoringViolation.objects.create(
                    session=session,
                    type=violation_type,
                    strike_number=strike_number,
                    comment=request.data.get('comment', '')
                )
            
            return Response(ProctoringViolationSerializer(violation).data, status=status.HTTP_201_CREATED)
        except AssessmentSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=status.HTTP_404_NOT_FOUND)

class AssessmentSessionCompleteView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def patch(self, request, session_id):
        try:
            session = AssessmentSession.objects.get(id=session_id, user=request.user)
            session.status = request.data.get('status', 'completed')
            session.score = request.data.get('score', 0.0)
            from django.utils import timezone
            session.completed_at = timezone.now()
            session.save()
            
            # Also create a permanent Result record
            # Handle AI customized tests without a formal assessment backing it
            assessment_obj = session.assessment
            if not assessment_obj:
                # Fallback for custom or missing assessment links
                assessment_obj, _ = Assessment.objects.get_or_create(
                    title="AI Personalized Skills Assessment",
                    defaults={
                        "description": "Adaptive test generated from uploaded CV",
                        "category": "AI Dynamic",
                        "difficulty": "Adaptive",
                        "estimated_time": 60
                    }
                )

            result_id = None
            # We always want a Result record now that assessment is nullable, 
            # but we've ensured we have a fallback assessment_obj above anyway.
            if True: 
                res = Result.objects.create(
                    user=request.user,
                    assessment=assessment_obj,
                    score=session.score,
                    coding_score=request.data.get('coding_score', None),
                    skill_breakdown=request.data.get('skill_breakdown', []),
                    correct_answers=request.data.get('correct_answers', 0),
                    incorrect_answers=request.data.get('incorrect_answers', 0),
                    total_questions=request.data.get('total_questions', 0),
                    time_taken=request.data.get('time_taken', 0),
                    ai_suggestions=request.data.get('ai_suggestions', []),
                    proctoring_summary=request.data.get('proctoring_summary', {})
                )
                result_id = res.id
            
            response_data = AssessmentSessionSerializer(session).data
            
            # Check if there was an active assignment and mark it complete
            try:
                if assessment_obj:
                    assignment = TestAssignment.objects.filter(candidate=request.user, assessment=assessment_obj, status='pending').first()
                    if assignment:
                        assignment.status = 'completed'
                        assignment.save()
            except Exception:
                pass
            
            if result_id:
                res_obj = Result.objects.get(id=result_id)
                response_data['result'] = ResultSerializer(res_obj).data
            
            # Auto-trigger Learning Roadmap Analysis (in background to avoid blocking)
            if result_id:
                user_obj = request.user
                def run_analysis():
                    try:
                        print(f"Starting background roadmap analysis for test result {result_id}")
                        services.generate_skill_gap_analysis(user_obj, 'skill_test', result_id)
                        print(f"Successfully completed roadmap analysis for test result {result_id}")
                    except Exception as e:
                        print(f"Error auto-generating roadmap: {e}")
                threading.Thread(target=run_analysis, daemon=True).start()
                    
            return Response(response_data)
        except AssessmentSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=status.HTTP_404_NOT_FOUND)

class CVTestGenerateView(views.APIView):
    permission_classes = (IsAuthenticated,) # Standardizing to IsAuthenticated
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        analysis_id = request.data.get('analysis_id')

        # 1. Handle Existing Analysis
        if analysis_id:
            try:
                cv_analysis = CVAnalysis.objects.get(id=analysis_id, user=request.user)
                return Response({
                    "id": cv_analysis.id,
                    "questions": cv_analysis.generated_questions
                }, status=status.HTTP_200_OK)
            except CVAnalysis.DoesNotExist:
                return Response({"error": "Analysis record not found"}, status=status.HTTP_404_NOT_FOUND)

        # 2. Handle New Upload
        if 'cv_file' not in request.FILES:
            return Response({"error": "No CV file provided"}, status=status.HTTP_400_BAD_REQUEST)
            
        api_key = request.data.get('api_key') or services.get_groq_api_key()
        if not api_key:
            return Response({"error": "No Groq API key provided"}, status=status.HTTP_400_BAD_REQUEST)

        file_obj = request.FILES['cv_file']
        
        try:
            cv_text = services.extract_text_from_file(file_obj)
            
            if not services.validate_cv_content(cv_text):
                return Response({"error": "Upload CV only. We couldn't find typical resume sections like Experience, Skills, or Education."}, status=status.HTTP_400_BAD_REQUEST)

            if len(cv_text.strip()) < 50:
                return Response({"error": "Could not extract enough text from the provided file."}, status=status.HTTP_400_BAD_REQUEST)
                
            json_response_str = services.generate_mcqs_from_cv(cv_text, api_key)
            
            try:
                # Robust extraction: find the first { and last }
                match = re.search(r'\{.*\}', json_response_str, re.DOTALL)
                if not match:
                    print(f"DEBUG: No JSON object found in AI response: {json_response_str}")
                    return Response({"error": "AI response did not contain a valid JSON object."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                cleaned_str = match.group(0)
                    
                questions_data = json.loads(cleaned_str)
                
                # PERSISTENCE: Save the analysis result to the database
                # Safely handle if AI returns an array directly instead of a dict
                is_dict = isinstance(questions_data, dict)
                skills = questions_data.get('skills', []) if is_dict else []
                # Store the entire raw questions_data to preserve coding_problems and metadata
                
                cv_analysis = CVAnalysis.objects.create(
                    user=request.user,
                    file_name=file_obj.name,
                    extracted_text=cv_text,
                    skills=skills,
                    generated_questions=questions_data
                )
                
                return Response({
                    "id": cv_analysis.id,
                    "questions": questions_data
                }, status=status.HTTP_200_OK)
                
            except json.JSONDecodeError as e:
                print(f"DEBUG: JSONDecodeError: {e}")
                print(f"DEBUG: Problematic string: {cleaned_str if 'cleaned_str' in locals() else 'N/A'}")
                return Response({"error": "Failed to parse AI-generated questions. Please try again."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class CodeforcesProblemFetchView(views.APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = services.fetch_codeforces_problems()
        if data.get('status') == 'OK':
            problems = data['result']['problems']
            if not problems:
                return Response({"error": "No problems found"}, status=status.HTTP_404_NOT_FOUND)
            
            # Select a random problem
            problem = random.choice(problems)
            
            # Formulate valid challenge data
            is_fallback = data.get('is_fallback', False)
            
            challenge_data = {
                "title": problem.get('name'),
                "difficulty": f"Rating: {problem.get('rating', 'N/A')}",
                "points": problem.get('points', 100),
                "tags": problem.get('tags', []),
                "contestId": problem.get('contestId'),
                "index": problem.get('index'),
                "url": problem.get('url') or f"https://codeforces.com/contest/{problem.get('contestId')}/problem/{problem.get('index')}",
                "description": problem.get('description') or f"This is a real-time challenge from Codeforces (Contest {problem.get('contestId')}, Problem {problem.get('index')}). Tags: {', '.join(problem.get('tags', []))}. Please solve it in the editor and verify your logic against Codeforces constraints.",
                "is_fallback": is_fallback
            }
            return Response(challenge_data)
        
        return Response({"error": data.get('comment', 'Failed to fetch from Codeforces')}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
class CodingSubmissionView(views.APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        serializer = CodingSubmissionSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(user=request.user)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ResultListView(generics.ListAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = ResultSerializer

    def get_queryset(self):
        return Result.objects.filter(user=self.request.user).order_by('-completed_at')

class ResultDetailView(generics.RetrieveAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = ResultSerializer

    def get_object(self):
        return Result.objects.get(id=self.kwargs['pk'], user=self.request.user)

class AIInterviewStartView(views.APIView):
    permission_classes = (IsAuthenticated,)
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request):
        assessment_id = request.data.get('assessment_id')
        analysis_id = request.data.get('analysis_id')
        cv_text = None
        file_name = None
        assessment_obj = None
        interview_context = ""

        # Handle Assigned Assessment
        if assessment_id:
            try:
                assessment_obj = Assessment.objects.get(id=assessment_id)
                # If it's an assigned interview, we prioritize its category (hr/technical)
                category = assessment_obj.interview_category or 'technical'
                interview_context = f"This is a {category.upper()} interview."
            except Assessment.DoesNotExist:
                return Response({"error": "Assigned Assessment not found"}, status=status.HTTP_404_NOT_FOUND)

        if analysis_id:
            try:
                analysis = CVAnalysis.objects.get(id=analysis_id, user=request.user)
                cv_text = analysis.extracted_text
                file_name = analysis.file_name
            except CVAnalysis.DoesNotExist:
                return Response({"error": "CV Analysis not found"}, status=status.HTTP_404_NOT_FOUND)
        
        # Fallback to user's latest CV if this is an assigned interview and no CV was uploaded
        if not cv_text and assessment_id:
            latest_analysis = CVAnalysis.objects.filter(user=request.user).order_by('-created_at').first()
            if latest_analysis:
                cv_text = latest_analysis.extracted_text
                file_name = latest_analysis.file_name
        
        if not cv_text:
            if 'cv_file' not in request.FILES:
                # If it's an assigned interview, we don't strictly NEED a CV file
                if not assessment_obj:
                    return Response({"error": "No CV file provided and no existing analysis found."}, status=status.HTTP_400_BAD_REQUEST)
                # Form a fallback profile based on the assessment title
                cv_text = f"Candidate Profile for {assessment_obj.title} ({assessment_obj.interview_category or 'general interview'})"
                file_name = "Assigned Profile"
            else:
                file_obj = request.FILES['cv_file']
                file_name = file_obj.name
                try:
                    cv_text = services.extract_text_from_file(file_obj)
                except Exception as e:
                    return Response({"error": f"Failed to extract text: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)

        api_key = request.data.get('api_key') or services.get_groq_api_key()
        if not api_key:
            return Response({"error": "Groq API key not found"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            if not cv_text or not services.validate_cv_content(cv_text):
                # If it's an assigned interview, we can fallback to a generic prompt if CV fails
                if not assessment_obj:
                    return Response({"error": "Invalid CV content or extraction failed"}, status=status.HTTP_400_BAD_REQUEST)
                cv_text = f"Candidate Profile for {assessment_obj.title}"

            # Pass the recruiter context to the generator
            interview_data = services.generate_ai_interview_questions(cv_text, api_key, context=interview_context)
            if not interview_data:
                return Response({"error": "Failed to generate interview questions"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            session = AIInterviewSession.objects.create(
                user=request.user,
                assessment=assessment_obj,
                file_name=file_name or "Assigned Assessment",
                extracted_text=cv_text,
                detected_skills=interview_data.get('skills', []),
                qa_history=[{
                    "question": interview_data.get('first_question'),
                    "topic": interview_data.get('topic'),
                    "difficulty": interview_data.get('difficulty', 'Easy'),
                    "timestamp": timezone.now().isoformat()
                }],
                status='in-progress'
            )

            return Response({
                "session_id": session.id,
                "first_question": interview_data.get('first_question'),
                "topic": interview_data.get('topic'),
                "skills": interview_data.get('skills'),
                "difficulty": interview_data.get('difficulty', 'Medium')
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AIInterviewNextQuestionView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request):
        session_id = request.data.get('session_id')
        last_answer = request.data.get('answer')
        
        if not session_id or not last_answer:
            return Response({"error": "Session ID and answer are required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            session = AIInterviewSession.objects.get(id=session_id, user=request.user)
            api_key = services.get_groq_api_key()

            # Update history with the last answer
            history = session.qa_history
            if history:
                history[-1]['answer'] = last_answer

            # Generate next question
            next_data = services.evaluate_and_generate_next_question(history, api_key)
            if not next_data:
                return Response({"error": "Failed to generate next question"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Append next question to history
            history.append({
                "question": next_data.get('next_question'),
                "topic": next_data.get('topic'),
                "evaluation_of_last_answer": next_data.get('evaluation_of_last_answer'),
                "timestamp": services.timezone.now().isoformat() if hasattr(services, 'timezone') else None
            })
            
            session.qa_history = history
            session.save()

            return Response({
                "next_question": next_data.get('next_question'),
                "topic": next_data.get('topic'),
                "evaluation": next_data.get('evaluation_of_last_answer'),
                "difficulty": next_data.get('difficulty', 'Medium')
            })

        except AIInterviewSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AIInterviewSubmitView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request):
        session_id = request.data.get('session_id')
        last_answer = request.data.get('answer')

        try:
            session = AIInterviewSession.objects.get(id=session_id, user=request.user)
            api_key = services.get_groq_api_key()

            history = session.qa_history
            if last_answer and history:
                history[-1]['answer'] = last_answer

            report_data = services.generate_interview_report(history, api_key)
            if not report_data:
                return Response({"error": "Failed to generate interview report"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            from django.utils import timezone
            session.overall_score = report_data.get('overall_score', 0)
            session.communication_score = report_data.get('communication_score', 0)
            session.technical_score = report_data.get('technical_score', 0)
            session.confidence_score = report_data.get('confidence_score', 0)
            session.answer_quality_score = report_data.get('answer_quality_score', 0)
            session.problem_solving_score = report_data.get('problem_solving_score', 0)
            session.feedback_summary = report_data.get('feedback_summary', '')
            session.strengths = report_data.get('strengths', [])
            session.improvements = report_data.get('improvements', [])
            session.status = 'completed'
            session.completed_at = timezone.now()
            session.qa_history = history
            session.save()

            # Create a Result record for the Dashboard
            # Use session assessment if available (for recruiter assigned interviews)
            assessment_obj = session.assessment
            if not assessment_obj:
                assessment_obj, _ = Assessment.objects.get_or_create(
                    title="AI Interactive Interview",
                    defaults={
                        "description": "Voice-based interactive technical interview.",
                        "category": "AI Interview",
                        "difficulty": "Adaptive",
                        "estimated_time": 15
                    }
                )

            # Mark any pending assignments as completed
            from .models import TestAssignment
            TestAssignment.objects.filter(candidate=request.user, assessment=assessment_obj, status='pending').update(status='completed')

            # Build comprehensive skill breakdown including technical skills
            skill_breakdown = [
                {"skill": "Technical Knowledge", "score": session.technical_score * 10},
                {"skill": "Communication", "score": session.communication_score * 10},
                {"skill": "Confidence", "score": session.confidence_score * 10},
                {"skill": "Answer Quality", "score": session.answer_quality_score * 10},
                {"skill": "Problem Solving", "score": session.problem_solving_score * 10}
            ]
            # Add specific detected skills for Career Engine matching
            for skill in session.detected_skills:
                skill_breakdown.append({"skill": skill, "score": session.technical_score * 10})

            Result.objects.create(
                user=request.user,
                assessment=assessment_obj,
                score=session.overall_score,
                ai_suggestions=report_data.get('strengths', []) + report_data.get('improvements', []),
                skill_breakdown=skill_breakdown,
                completed_at=session.completed_at
            )

            # Auto-trigger Learning Roadmap Analysis (in background to avoid blocking)
            user_obj = request.user
            session_id = session.id
            def run_interview_analysis():
                try:
                    print(f"Starting background skill gap analysis for interview session {session_id}")
                    services.generate_skill_gap_analysis(user_obj, 'ai_interview', session_id)
                    print(f"Successfully completed background analysis for session {session_id}")
                except Exception as e:
                    print(f"Error auto-generating interview roadmap: {e}")
            threading.Thread(target=run_interview_analysis, daemon=True).start()

            return Response(report_data)

        except AIInterviewSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AIInterviewCVStatusView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        latest_analysis = CVAnalysis.objects.filter(user=request.user).order_by('-created_at').first()
        if latest_analysis:
            return Response({
                "exists": True,
                "file_name": latest_analysis.file_name,
                "skills": latest_analysis.skills,
                "id": latest_analysis.id,
                "created_at": latest_analysis.created_at
            })
        
        # Check if they have a resume in profile but not analyzed
        profile = getattr(request.user, 'profile', None)
        if profile and profile.resume:
            return Response({
                "exists": False,
                "has_raw_file": True,
                "file_name": os.path.basename(profile.resume.name)
            })

        return Response({"exists": False, "has_raw_file": False})

class SkillGapAnalysisListView(generics.ListAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = SkillGapAnalysisSerializer

    def get_queryset(self):
        return SkillGapAnalysis.objects.filter(user=self.request.user)

class SkillGapAnalysisDetailView(generics.RetrieveDestroyAPIView):
    permission_classes = (IsAuthenticated,)
    serializer_class = SkillGapAnalysisSerializer

    def get_queryset(self):
        return SkillGapAnalysis.objects.filter(user=self.request.user)

class SkillGapAnalysisGenerateView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request):
        source_type = request.data.get('source_type')
        source_id = request.data.get('source_id')

        if not source_type or not source_id:
            return Response({"error": "source_type and source_id are required"}, status=status.HTTP_400_BAD_REQUEST)

        # check if already exists to avoid redundant AI calls
        existing = SkillGapAnalysis.objects.filter(
            user=request.user, 
            source_type=source_type
        )
        if source_type == 'skill_test':
            existing = existing.filter(assessment_result_id=source_id)
        else:
            existing = existing.filter(interview_session_id=source_id)
        
        if existing.exists():
            return Response(SkillGapAnalysisSerializer(existing.first()).data)

        analysis = services.generate_skill_gap_analysis(request.user, source_type, source_id)
        if not analysis:
            return Response({"error": "Failed to generate skill gap analysis"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response(SkillGapAnalysisSerializer(analysis).data, status=status.HTTP_201_CREATED)

class LeaderboardView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        from django.db.models import Avg, FloatField
        from django.db.models.functions import Coalesce
        from .models import User, Result, AIInterviewSession

        # Aggregate average test scores
        test_scores = Result.objects.values('user_id').annotate(
            avg_test=Avg('score')
        )
        test_map = {item['user_id']: item['avg_test'] for item in test_scores}

        # Aggregate average interview scores
        interview_scores = AIInterviewSession.objects.filter(status='completed').values('user_id').annotate(
            avg_interview=Avg('overall_score')
        )
        interview_map = {item['user_id']: item['avg_interview'] for item in interview_scores}

        leaderboard = []
        users = User.objects.filter(is_active=True, role='student')
        
        for u in users:
            ts = test_map.get(u.id)
            ins = interview_map.get(u.id)
            
            # Scoring Logic: Average of test and interview averages
            # If only one exists, use that.
            scores = []
            if ts is not None: scores.append(ts)
            if ins is not None: scores.append(ins)
            
            final_score = sum(scores) / len(scores) if scores else 0
            
            if final_score > 0 or u == request.user: # Include user even if 0
                leaderboard.append({
                    "id": u.id,
                    "name": f"{u.first_name} {u.last_name[:1]}." if u.first_name and u.last_name else u.username,
                    "score": round(final_score, 1),
                    "isUser": u == request.user
                })
        
        # Sort by score descending
        leaderboard = sorted(leaderboard, key=lambda x: x['score'], reverse=True)
        
        # Add rank
        for i, item in enumerate(leaderboard):
            item['rank'] = i + 1
            
        return Response(leaderboard[:10])

class LatestAnalysisView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        from .models import SkillGapAnalysis
        from .serializers import SkillGapAnalysisSerializer
        
        latest = SkillGapAnalysis.objects.filter(user=request.user).order_by('-created_at').first()
        if not latest:
            return Response({"exists": False})
            
        return Response({
            "exists": True,
            "data": SkillGapAnalysisSerializer(latest).data
        })

class RecruiterDashboardAnalyticsView(views.APIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        from django.db.models import Avg, Count
        from .models import User, Result, AIInterviewSession
        
        # 1. Core Metrics
        total_candidates = User.objects.filter(role='student').count()
        total_assessments = Result.objects.count()
        
        avg_test_score = Result.objects.aggregate(avg=Avg('score'))['avg'] or 0
        avg_interview_score = AIInterviewSession.objects.filter(status='completed').aggregate(avg=Avg('overall_score'))['avg'] or 0
        
        # 2. Pipeline Data (Score brackets)
        # Bronze: 0-40, Silver: 41-75, Gold: 76-90, Platinum: 91-100
        pipeline = [
            {"label": "Bronze (0-40%)", "value": Result.objects.filter(score__lte=40).count(), "colorClass": "error"},
            {"label": "Silver (41-75%)", "value": Result.objects.filter(score__gt=40, score__lte=75).count(), "colorClass": "warning"},
            {"label": "Gold (76-90%)", "value": Result.objects.filter(score__gt=75, score__lte=90).count(), "colorClass": "info"},
            {"label": "Platinum (91-100%)", "value": Result.objects.filter(score__gt=90).count(), "colorClass": "success"},
        ]
        
        # 3. Recent Activity (Latest 5 events)
        recent_results = Result.objects.select_related('user', 'assessment').order_by('-completed_at')[:5]
        recent_interviews = AIInterviewSession.objects.filter(status='completed').select_related('user').order_by('-completed_at')[:5]
        
        activity_feed = []
        for r in recent_results:
            activity_feed.append({
                "id": f"res_{r.id}",
                "type": "assessment",
                "user": f"{r.user.first_name} {r.user.last_name[:1]}.",
                "content": f"completed {r.assessment.title}",
                "time": r.completed_at.strftime("%Y-%m-%d %H:%M"),
                "score": f"{round(r.score)}%"
            })
            
        for i in recent_interviews:
            activity_feed.append({
                "id": f"int_{i.id}",
                "type": "interview",
                "user": f"{i.user.first_name} {i.user.last_name[:1]}.",
                "content": "completed AI Interview",
                "time": i.created_at.strftime("%Y-%m-%d %H:%M"),
                "score": f"{i.overall_score}%"
            })
            
        # Sort combined activity and take top 5
        activity_feed = sorted(activity_feed, key=lambda x: x['time'], reverse=True)[:5]
        
        return Response({
            "metrics": {
                "total_candidates": total_candidates,
                "total_assessments": total_assessments,
                "avg_score": round((avg_test_score + avg_interview_score) / 2 if avg_test_score and avg_interview_score else max(avg_test_score, avg_interview_score), 1),
                "active_interviews": AIInterviewSession.objects.filter(status='in-progress').count()
            },
            "pipeline": pipeline,
            "activity": activity_feed
        })

class StudentTestAssignmentView(generics.ListAPIView):
    """
    API view for students to fetch assigned exams by a recruiter.
    """
    permission_classes = [IsAuthenticated]
    serializer_class = TestAssignmentSerializer

    def get_queryset(self):
        return TestAssignment.objects.filter(candidate=self.request.user, status='pending').order_by('-created_at')

class VerifyInvitationView(views.APIView):
    permission_classes = [AllowAny]
    
    def get(self, request, token):
        try:
            invitation = TestInvitation.objects.get(token=token)
            
            # Check deadline
            if invitation.deadline and invitation.deadline < timezone.now():
                if invitation.status != 'expired':
                    invitation.status = 'expired'
                    invitation.save()
                return Response({'error': 'This test invitation has expired.'}, status=status.HTTP_400_BAD_REQUEST)
                
            if invitation.status == 'completed':
                return Response({'error': 'This test invitation has already been completed.'}, status=status.HTTP_400_BAD_REQUEST)
            
            # Mark as opened if pending
            if invitation.status == 'pending':
                invitation.status = 'opened'
                invitation.save()
                
            serializer = TestInvitationSerializer(invitation)
            return Response({
                'invitation': serializer.data,
                'recruiter_name': invitation.recruiter.get_full_name() or invitation.recruiter.username
            })
            
        except TestInvitation.DoesNotExist:
            return Response({'error': 'Invalid or expired invitation link.'}, status=status.HTTP_404_NOT_FOUND)

class CompleteInvitationView(views.APIView):
    permission_classes = [AllowAny]
    
    def post(self, request, token):
        try:
            invitation = TestInvitation.objects.get(token=token)
            invitation.status = 'completed'
            invitation.save()
            return Response({'message': 'Status updated to completed.'})
        except TestInvitation.DoesNotExist:
            return Response({'error': 'Invalid token.'}, status=status.HTTP_404_NOT_FOUND)
